"""
Document converter module using Docling for high-quality PDF/DOCX to Markdown conversion.
"""

import os
import json
from pathlib import Path
from datetime import datetime
from typing import Optional, Tuple
from dataclasses import dataclass, asdict

try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import PdfFormatOption
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
    print("Warning: Docling not installed. Run: pip install docling")


@dataclass
class DocumentMetadata:
    """Metadata extracted from and about the document."""
    filename: str
    original_path: str
    file_extension: str
    file_size_bytes: int
    date_created: Optional[str] = None  # File creation date
    date_modified: Optional[str] = None  # File modification date
    date_processed: str = ""  # When we processed it
    title: Optional[str] = None  # Extracted from document
    author: Optional[str] = None  # Extracted from document
    page_count: Optional[int] = None
    has_tables: bool = False
    has_images: bool = False
    
    def __post_init__(self):
        if not self.date_processed:
            self.date_processed = datetime.now().isoformat()
    
    def to_dict(self) -> dict:
        return asdict(self)
    
    def to_json(self) -> str:
        return json.dumps(self.to_dict(), indent=2)


class DoclingConverter:
    """
    Converts PDF and DOCX files to high-quality Markdown using Docling.
    Preserves tables, formatting, and extracts metadata.
    """
    
    def __init__(self):
        if not DOCLING_AVAILABLE:
            raise ImportError(
                "Docling is not installed. Install with: pip install docling"
            )
        
        # Configure pipeline for better table handling
        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = False  # Disable OCR for faster processing
        pipeline_options.do_table_structure = True  # Enable table structure recognition
        
        self.converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
    
    def convert_file(self, filepath: Path) -> Tuple[str, DocumentMetadata]:
        """
        Convert a document to Markdown.
        
        Args:
            filepath: Path to the PDF or DOCX file
            
        Returns:
            Tuple of (markdown_content, metadata)
        """
        filepath = Path(filepath)
        
        if not filepath.exists():
            raise FileNotFoundError(f"File not found: {filepath}")
        
        # Gather file metadata
        metadata = self._extract_metadata(filepath)
        
        # Convert document
        result = self.converter.convert(str(filepath))
        
        # Export to Markdown
        markdown_content = result.document.export_to_markdown()
        
        # Update metadata with document info
        if hasattr(result.document, 'metadata'):
            doc_meta = result.document.metadata
            if hasattr(doc_meta, 'title') and doc_meta.title:
                metadata.title = doc_meta.title
            if hasattr(doc_meta, 'author') and doc_meta.author:
                metadata.author = doc_meta.author
        
        # Check for tables in the output
        metadata.has_tables = '|' in markdown_content and '---' in markdown_content
        
        # Estimate page count from content structure
        if hasattr(result.document, 'pages'):
            metadata.page_count = len(result.document.pages)
        
        return markdown_content, metadata
    
    def _extract_metadata(self, filepath: Path) -> DocumentMetadata:
        """Extract file system metadata from the document."""
        stat = filepath.stat()
        
        # Get file dates
        date_created = None
        date_modified = None
        
        try:
            # Creation time (may not be available on all systems)
            if hasattr(stat, 'st_birthtime'):
                date_created = datetime.fromtimestamp(stat.st_birthtime).isoformat()
            else:
                # Fall back to ctime (metadata change time on Unix)
                date_created = datetime.fromtimestamp(stat.st_ctime).isoformat()
            
            date_modified = datetime.fromtimestamp(stat.st_mtime).isoformat()
        except Exception:
            pass
        
        return DocumentMetadata(
            filename=filepath.name,
            original_path=str(filepath.absolute()),
            file_extension=filepath.suffix.lower(),
            file_size_bytes=stat.st_size,
            date_created=date_created,
            date_modified=date_modified,
        )


class FallbackConverter:
    """
    Fallback converter for when Docling is not available.
    Uses basic extraction methods.
    """
    
    def convert_file(self, filepath: Path) -> Tuple[str, DocumentMetadata]:
        """Basic conversion fallback."""
        filepath = Path(filepath)
        
        if not filepath.exists():
            raise FileNotFoundError(f"File not found: {filepath}")
        
        stat = filepath.stat()
        metadata = DocumentMetadata(
            filename=filepath.name,
            original_path=str(filepath.absolute()),
            file_extension=filepath.suffix.lower(),
            file_size_bytes=stat.st_size,
        )
        
        # Try basic extraction based on file type
        if filepath.suffix.lower() == '.pdf':
            markdown = self._extract_pdf_basic(filepath)
        elif filepath.suffix.lower() in ['.docx', '.doc']:
            markdown = self._extract_docx_basic(filepath)
        else:
            raise ValueError(f"Unsupported file type: {filepath.suffix}")
        
        return markdown, metadata
    
    def _extract_pdf_basic(self, filepath: Path) -> str:
        """Basic PDF text extraction using PyPDF2 or pdfplumber."""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                    
                    # Try to extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            md_table = self._table_to_markdown(table)
                            text_parts.append(md_table)
            
            return '\n\n'.join(text_parts)
        except ImportError:
            pass
        
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(filepath)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return '\n\n'.join(text_parts)
        except ImportError:
            raise ImportError("No PDF library available. Install pdfplumber or pypdf.")
    
    def _extract_docx_basic(self, filepath: Path) -> str:
        """Basic DOCX text extraction."""
        try:
            from docx import Document
            
            doc = Document(filepath)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # Check if it's a heading
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                        paragraphs.append(f"{'#' * level} {para.text}")
                    else:
                        paragraphs.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                md_table = self._docx_table_to_markdown(table)
                paragraphs.append(md_table)
            
            return '\n\n'.join(paragraphs)
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    def _table_to_markdown(self, table: list) -> str:
        """Convert a table (list of lists) to Markdown."""
        if not table or not table[0]:
            return ""
        
        lines = []
        # Header row
        header = table[0]
        lines.append('| ' + ' | '.join(str(cell or '') for cell in header) + ' |')
        lines.append('| ' + ' | '.join('---' for _ in header) + ' |')
        
        # Data rows
        for row in table[1:]:
            lines.append('| ' + ' | '.join(str(cell or '') for cell in row) + ' |')
        
        return '\n'.join(lines)
    
    def _docx_table_to_markdown(self, table) -> str:
        """Convert a python-docx table to Markdown."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        
        return self._table_to_markdown(rows)


def get_converter():
    """Get the best available converter."""
    if DOCLING_AVAILABLE:
        return DoclingConverter()
    else:
        print("Warning: Using fallback converter. Install docling for better results.")
        return FallbackConverter()
